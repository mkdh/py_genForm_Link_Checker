import docx
from datetime import datetime

# Returns a datetime object containing the local date and time
global_dateTimeObj = datetime.now()
global_timestamp = str(global_dateTimeObj.year) + str(global_dateTimeObj.month).zfill(2) + str(global_dateTimeObj.day).zfill(2) + str(global_dateTimeObj.hour).zfill(2) + str(global_dateTimeObj.minute).zfill(2) + str(global_dateTimeObj.second).zfill(2)# + str(global_dateTimeObj.microsecond)

# --- Load file
import os
def get_file_by_type(my_type, my_folder = 'my_data'):
    for file in os.listdir(my_folder):
        if file.endswith("." + my_type):
            return os.path.join(my_folder, file)

    return ""

def get_docx():
    return get_file_by_type("docx")

document = docx.Document(get_docx()) #load file
# --- end load

tables = document.tables #获取文件中的表格集


def replaceText(document, search, replace):
    for table in document.tables:
        for row in table.rows:
            for paragraph in row.cells:
                if search in paragraph.text:
                    paragraph.text = replace

def createNewWord(fileName, table):    
    for i in range(0, len(table.rows)):#從表格第二行開始循環讀取表格數據
        replaceText(newDoc, 'Text_Fill_Date', str(global_dateTimeObj.year - 1911) + '.' + str(global_dateTimeObj.month) + '.' + str(global_dateTimeObj.day))
        idNum = table.cell(i,0).text #
        
        if idNum == '主網站名稱':
            replaceText(newDoc, 'Text_Authority_Unit', table.cell(i,1).text)
            
        if idNum == '單元名稱':
            replaceText(newDoc, 'Text_Page_Location', table.cell(i,1).text)
            
        if idNum == '連結網址':
            replaceText(newDoc, 'Text_Unit_Link', table.cell(i,1).text)


    newDoc.save(fileName)

def findValue(table, keyword):
    retValue = "Unknown"
    for i in range(0, len(table.rows)):#從表格第二行開始循環讀取表格數據
        idNum = table.cell(i,0).text #
        if idNum == keyword:
            retValue = table.cell(i,1).text
    return retValue


i = 0

for myTable in document.tables:  
    newDoc = docx.Document(get_file_by_type('docx', 'my_sys'))
    i = i + 1
    createNewWord(global_timestamp + "_" + str(i) + "_" + findValue(myTable, '主網站名稱') + ".docx",myTable)


	
# reference:
# image https://stackoverflow.com/questions/27691678/finding-image-present-docx-file-using-python?noredirect=1&lq=1
# https://gifguide2code.com/2017/12/11/python-how-to-deal-with-tables-in-word/


